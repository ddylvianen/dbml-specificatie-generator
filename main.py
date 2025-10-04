from docx import Document
from docx.shared import Pt
from pydbml import PyDBML

db = PyDBML.parse_file('database.dbml')
tables = db.tables

doc = Document()

for table in tables:
    table_doc = doc.add_table(rows=1, cols=5)
    table_doc.style = 'Table Grid'
    table_doc.borders = True
    
    
    
    hdr_cells = table_doc.rows[0].cells
    hdr_cells[0].text = table.name
    hdr_cells[0].merge(hdr_cells[4])

    hdr_cells = table_doc.add_row().cells
    hdr_cells[0].text = "Kolomnaam"
    hdr_cells[1].text = "Datatype"
    hdr_cells[2].text = "Lengte"
    hdr_cells[3].text = "Nullable"
    hdr_cells[4].text = "Opmerking"

    for column in table.columns:
        hdr_cells = table_doc.add_row().cells
        hdr_cells[0].text = column.name
        hdr_cells[1].text = column.type.split("(", 1)[0] if "(" in column.type else column.type
        hdr_cells[2].text = column.type.split("(", 1)[1].rstrip(")") if "(" in column.type else ""
        hdr_cells[3].text = { True: "NOT NULL", False: "NULL" }[column.not_null]
        hdr_cells[4].text = f"PK" if column.pk else ""
    doc.add_paragraph("\n")



doc.save('specificatie.docx')